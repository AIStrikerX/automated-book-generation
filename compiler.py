"""
Compiler Module
Compiles chapters into final book document (.docx, .txt, PDF)
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict


class BookCompiler:
    """Handles compilation of chapters into final book format"""
    
    def __init__(self, output_dir: str = "output"):
        """
        Initialize compiler
        
        Args:
            output_dir: Directory to save compiled books
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def compile_to_docx(self, title: str, chapters: List[Dict], 
                       book_id: str, include_metadata: bool = True) -> str:
        """
        Compile chapters into a Word document
        
        Args:
            title: Book title
            chapters: List of chapter dictionaries with 'chapter_number', 'content', etc.
            book_id: Unique book identifier
            include_metadata: Whether to include generation metadata
            
        Returns:
            Path to the generated .docx file
        """
        document = Document()
        
        # Set document margins
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title page
        title_para = document.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if include_metadata:
            date_para = document.add_paragraph(f"\nGenerated on: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(10)
        
        document.add_page_break()
        
        # Add table of contents placeholder
        document.add_heading("Table of Contents", 1)
        for chapter in chapters:
            ch_title = chapter.get('chapter_title') or chapter.get('title') or f'Chapter {chapter["chapter_number"]}'
            toc_entry = document.add_paragraph(
                f"Chapter {chapter['chapter_number']}: {ch_title}"
            )
            toc_entry.style = 'List Number'
        
        document.add_page_break()
        
        # Add chapters
        for chapter in sorted(chapters, key=lambda x: x['chapter_number']):
            # Chapter heading
            chapter_heading = document.add_heading(
                f"Chapter {chapter['chapter_number']}", 
                level=1
            )
            
            # Add chapter title if available
            ch_title = chapter.get('chapter_title') or chapter.get('title', '')
            if ch_title:
                title_para = document.add_heading(ch_title, level=2)
            
            # Add chapter content
            content = chapter['content']
            
            # Split content into paragraphs and handle formatting
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                
                # Check if it's a heading (starts with ##, ###, etc.)
                if para_text.startswith('###'):
                    document.add_heading(para_text.replace('###', '').strip(), level=3)
                elif para_text.startswith('##'):
                    document.add_heading(para_text.replace('##', '').strip(), level=2)
                elif para_text.startswith('#'):
                    document.add_heading(para_text.replace('#', '').strip(), level=1)
                else:
                    # Regular paragraph
                    para = document.add_paragraph(para_text)
                    para.paragraph_format.line_spacing = 1.5
            
            # Add page break after each chapter except the last
            if chapter != chapters[-1]:
                document.add_page_break()
        
        # Save document
        filename = f"{self._sanitize_filename(title)}_{book_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        document.save(filepath)
        
        return filepath
    
    def compile_to_txt(self, title: str, chapters: List[Dict], 
                      book_id: str, include_metadata: bool = True) -> str:
        """
        Compile chapters into a plain text file
        
        Args:
            title: Book title
            chapters: List of chapter dictionaries
            book_id: Unique book identifier
            include_metadata: Whether to include generation metadata
            
        Returns:
            Path to the generated .txt file
        """
        content_lines = []
        
        # Add title
        content_lines.append("=" * 80)
        content_lines.append(title.center(80))
        content_lines.append("=" * 80)
        content_lines.append("")
        
        if include_metadata:
            content_lines.append(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
            content_lines.append("")
        
        content_lines.append("")
        
        # Add table of contents
        content_lines.append("TABLE OF CONTENTS")
        content_lines.append("-" * 80)
        for chapter in sorted(chapters, key=lambda x: x['chapter_number']):
            ch_title = chapter.get('chapter_title') or chapter.get('title') or f'Chapter {chapter["chapter_number"]}'
            content_lines.append(
                f"Chapter {chapter['chapter_number']}: {ch_title} "
            )
        content_lines.append("")
        content_lines.append("=" * 80)
        content_lines.append("")
        
        # Add chapters
        for chapter in sorted(chapters, key=lambda x: x['chapter_number']):
            content_lines.append("")
            content_lines.append("=" * 80)
            content_lines.append(f"CHAPTER {chapter['chapter_number']}")
            ch_title = chapter.get('chapter_title') or chapter.get('title', '')
            if ch_title:
                content_lines.append(ch_title)
            content_lines.append("=" * 80)
            content_lines.append("")
            content_lines.append(chapter['content'])
            content_lines.append("")
        
        # Save to file
        filename = f"{self._sanitize_filename(title)}_{book_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content_lines))
        
        return filepath
    
    def compile_both_formats(self, title: str, chapters: List[Dict], 
                           book_id: str) -> Dict[str, str]:
        """
        Compile chapters into .docx, .txt, and .pdf formats
        
        Returns:
            Dictionary with 'docx', 'txt', and 'pdf' file paths
        """
        docx_path = self.compile_to_docx(title, chapters, book_id)
        txt_path = self.compile_to_txt(title, chapters, book_id)
        pdf_path = self.compile_to_pdf(docx_path, title, book_id)
        
        return {
            'docx': docx_path,
            'txt': txt_path,
            'pdf': pdf_path
        }
    
    def compile_to_pdf(self, docx_path: str, title: str, book_id: str) -> str:
        """
        Convert .docx to .pdf format
        
        Args:
            docx_path: Path to the .docx file
            title: Book title
            book_id: Unique identifier for the book
            
        Returns:
            Path to the generated PDF file
        """
        from docx2pdf import convert
        import os
        
        # Generate PDF filename
        clean_title = self._sanitize_filename(title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"{clean_title}_{timestamp}.pdf"
        pdf_path = os.path.join(self.output_dir, pdf_filename)
        
        try:
            # Convert DOCX to PDF
            convert(docx_path, pdf_path)
            print(f"✅ PDF created: {pdf_path}")
            return pdf_path
        except Exception as e:
            print(f"⚠️  PDF conversion failed: {e}")
            print(f"   (DOCX file is still available at: {docx_path})")
            return None
    
    def _sanitize_filename(self, filename: str) -> str:
        """Remove invalid characters from filename"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        # Limit length
        return filename[:50]
